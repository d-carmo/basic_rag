from pathlib import Path

import docx as python_docx

from rag.loaders.base import BaseLoader, DocType, Document, DocumentMetadata


class DocxLoader(BaseLoader):
    @property
    def supported_extensions(self) -> frozenset[str]:
        return frozenset({".docx", ".doc"})

    def load(self, source: Path | str) -> list[Document]:
        path = Path(source)
        doc = python_docx.Document(str(path))

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)
        if not text.strip():
            return []

        return [
            Document(
                text=text,
                metadata=DocumentMetadata(
                    source_url=str(path.resolve()),
                    doc_type=DocType.DOCX,
                ),
            )
        ]
